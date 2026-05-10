import logging
import asyncio
from typing import List, Dict, Any
from PyPDF2 import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy.orm import Session
from fastapi import UploadFile
from app.config import settings
from app.exceptions import DocumentNotFoundError, DocumentProcessingError, UnsupportedFileTypeError, FileSizeLimitError
from app.models.document import Document
from app.utils.file_utils import save_upload, delete_upload, SUPPORTED_EXTENSIONS
from app.ai.embedder import embed_documents
from app.ai.vector_store import store_chunks, delete_document_chunks

logger = logging.getLogger(__name__)


def validate_file(file: UploadFile) -> None:
    # 1. Validate Extension
    import os
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(extension=ext, supported=SUPPORTED_EXTENSIONS)

    # 2. Validate Size
    # FastAPI SpooledTemporaryFile might not have accurate seek depending on read state,
    # but we can check the spool size or read bytes. We'll rely on the caller to pass bytes.
    pass


class DocumentService:
    def __init__(self, db: Session):
        self.db = db

    async def upload_document(self, file: UploadFile, user_id: str) -> Document:
        validate_file(file)

        file_bytes = await file.read()
        file_size  = len(file_bytes)

        if file_size > settings.MAX_FILE_SIZE_MB * 1024 * 1024:
            raise FileSizeLimitError(actual_bytes=file_size, max_bytes=settings.MAX_FILE_SIZE_MB * 1024 * 1024)

        # Save to disk
        file_path = save_upload(file_bytes, file.filename or "unknown", user_id)

        # Create DB record (status=processing)
        doc = Document(
            user_id=user_id,
            filename=file.filename,
            file_path=file_path,
            file_size=file_size,
            status="processing"
        )
        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)

        # Kick off background processing
        # In a real production app, this would use Celery/Redis.
        # For this exercise, we use asyncio.create_task.
        # We need a new DB session for the background task since the current one will close.
        from app.database import SessionLocal
        bg_db = SessionLocal()
        asyncio.create_task(self._process_document_background(doc.id, bg_db))

        return doc

    def list_documents(self, user_id: str) -> list[Document]:
        return self.db.query(Document).filter(Document.user_id == user_id).order_by(Document.created_at.desc()).all()

    def get_document(self, doc_id: str, user_id: str) -> Document:
        doc = self.db.query(Document).filter(Document.id == doc_id, Document.user_id == user_id).first()
        if not doc:
            raise DocumentNotFoundError(doc_id)
        return doc

    def delete_document(self, doc_id: str, user_id: str) -> None:
        doc = self.get_document(doc_id, user_id)
        
        # 1. Delete chunks from ChromaDB
        try:
            delete_document_chunks(user_id, doc_id)
        except Exception as e:
            logger.warning(f"Failed to delete Chroma chunks for {doc_id}: {e}")

        # 2. Delete file from disk
        delete_upload(doc.file_path)

        # 3. Delete from SQL
        self.db.delete(doc)
        self.db.commit()

    async def _process_document_background(self, doc_id: str, db: Session) -> None:
        """Background task to extract text, chunk, embed, and store in ChromaDB."""
        try:
            doc = db.query(Document).filter(Document.id == doc_id).first()
            if not doc:
                return

            # 1. Extract Text
            chunks_data = self._extract_text(doc.file_path)
            doc.page_count = len(set(c["page_number"] for c in chunks_data))

            # 2. Chunking
            splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            final_chunks = []
            for c in chunks_data:
                splits = splitter.split_text(c["text"])
                for s in splits:
                    final_chunks.append({
                        "text": s,
                        "page_number": c["page_number"]
                    })

            if not final_chunks:
                raise ValueError("No text could be extracted from the document.")

            # 3. Embed (batch)
            texts_to_embed = [c["text"] for c in final_chunks]
            embeddings = embed_documents(texts_to_embed)

            for chunk, emb in zip(final_chunks, embeddings):
                chunk["embedding"] = emb

            # 4. Store in ChromaDB
            store_chunks(
                user_id=doc.user_id,
                document_id=doc.id,
                filename=doc.filename,
                chunks=final_chunks
            )

            # 5. Update Status
            doc.status = "ready"
            db.commit()

        except Exception as e:
            logger.error(f"Failed to process document {doc_id}: {e}", exc_info=True)
            doc = db.query(Document).filter(Document.id == doc_id).first()
            if doc:
                doc.status = "error"
                db.commit()
        finally:
            db.close()

    def _extract_text(self, file_path: str) -> list[Dict[str, Any]]:
        """
        Extract text from supported files.
        Returns list of dicts: [{"text": str, "page_number": int}]
        """
        ext = file_path.lower().split('.')[-1]
        chunks = []

        if ext == 'pdf':
            reader = PdfReader(file_path)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    chunks.append({"text": text.strip(), "page_number": i + 1})
                    
        elif ext in ['txt', 'md']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
                if text.strip():
                    chunks.append({"text": text.strip(), "page_number": 1})
                    
        elif ext == 'docx':
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            text = '\n'.join(full_text)
            if text.strip():
                chunks.append({"text": text.strip(), "page_number": 1})
                
        else:
            raise ValueError(f"Extraction not implemented for .{ext}")

        return chunks
